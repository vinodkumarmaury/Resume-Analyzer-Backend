import os
import PyPDF2
import pdfplumber
from docx import Document
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from collections import Counter
import math
import numpy as np
from .models import ResumeAnalysis
from jobs.models import Job

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

def parse_resume(file_path):
    """Extract text and parse data from resume file"""
    
    extracted_text = ""
    
    if file_path.lower().endswith('.pdf'):
        extracted_text = extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        extracted_text = extract_text_from_docx(file_path)
    
    # Parse the extracted text
    parsed_data = parse_resume_text(extracted_text)
    
    return extracted_text, parsed_data

def extract_text_from_pdf(file_path):
    """Extract text from PDF file using both PyPDF2 and pdfplumber"""
    text = ""
    
    # Try pdfplumber first (better for complex layouts)
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                    
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")
    
    return text

def parse_resume_text(text):
    """Parse resume text to extract structured data"""
    
    parsed_data = {
        'contact_info': extract_contact_info(text),
        'skills': extract_skills_advanced(text),
        'experience': extract_experience(text),
        'education': extract_education(text),
        'sections': identify_sections(text),
        'keywords': extract_keywords_simple(text)
    }
    
    return parsed_data

def extract_contact_info(text):
    """Extract contact information from resume text"""
    contact_info = {}
    
    # Email regex
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails:
        contact_info['email'] = emails[0]
    
    # Phone regex (multiple formats)
    phone_patterns = [
        r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'(\+?\d{1,3}[-.\s]?)?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
        r'(\+?\d{1,3}[-.\s]?)?\d{10}'
    ]
    
    for pattern in phone_patterns:
        phones = re.findall(pattern, text)
        if phones:
            contact_info['phone'] = phones[0] if isinstance(phones[0], str) else ''.join(phones[0])
            break
    
    # LinkedIn URL
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'
    linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
    if linkedin:
        contact_info['linkedin'] = f"https://{linkedin[0]}"
    
    # GitHub URL
    github_pattern = r'github\.com/[\w-]+'
    github = re.findall(github_pattern, text, re.IGNORECASE)
    if github:
        contact_info['github'] = f"https://{github[0]}"
    
    return contact_info

def extract_skills_advanced(text):
    """Extract skills using NLTK and predefined skill lists"""
    
    # Comprehensive skill database
    technical_skills = [
        # Programming Languages
        'Python', 'JavaScript', 'Java', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust', 'Swift',
        'Kotlin', 'TypeScript', 'Scala', 'R', 'MATLAB', 'Perl', 'Shell', 'Bash',
        
        # Web Technologies
        'HTML', 'CSS', 'React', 'Angular', 'Vue.js', 'Node.js', 'Express.js', 'Django',
        'Flask', 'FastAPI', 'Spring', 'Laravel', 'Rails', 'ASP.NET', 'jQuery', 'Bootstrap',
        'Tailwind CSS', 'SASS', 'LESS', 'Webpack', 'Vite', 'Next.js', 'Nuxt.js',
        
        # Databases
        'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'SQLite', 'Oracle', 'SQL Server',
        'Cassandra', 'DynamoDB', 'Elasticsearch', 'Neo4j', 'Firebase',
        
        # Cloud & DevOps
        'AWS', 'Azure', 'Google Cloud', 'Docker', 'Kubernetes', 'Jenkins', 'GitLab CI',
        'GitHub Actions', 'Terraform', 'Ansible', 'Chef', 'Puppet', 'Vagrant',
        
        # Data Science & AI
        'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn',
        'Pandas', 'NumPy', 'Matplotlib', 'Seaborn', 'Jupyter', 'Apache Spark', 'Hadoop',
        'Tableau', 'Power BI', 'D3.js', 'OpenCV', 'NLTK', 'spaCy',
        
        # Mobile Development
        'iOS', 'Android', 'React Native', 'Flutter', 'Xamarin', 'Ionic',
        
        # Tools & Others
        'Git', 'SVN', 'Jira', 'Confluence', 'Slack', 'Trello', 'Figma', 'Adobe Creative Suite',
        'Photoshop', 'Illustrator', 'InDesign', 'Sketch', 'InVision', 'Zeplin'
    ]
    
    # Soft skills
    soft_skills = [
        'Leadership', 'Communication', 'Teamwork', 'Problem Solving', 'Critical Thinking',
        'Project Management', 'Time Management', 'Adaptability', 'Creativity', 'Innovation',
        'Analytical Skills', 'Attention to Detail', 'Customer Service', 'Negotiation',
        'Presentation Skills', 'Public Speaking', 'Mentoring', 'Coaching'
    ]
    
    all_skills = technical_skills + soft_skills
    found_skills = []
    text_lower = text.lower()
    
    # Direct skill matching
    for skill in all_skills:
        if skill.lower() in text_lower:
            found_skills.append(skill)
    
    # Use NLTK for additional entity-like extraction
    try:
        tokens = word_tokenize(text)
        # Look for capitalized words that might be technologies
        for token in tokens:
            if token.istitle() and len(token) > 2:
                # Check if it's similar to known technologies
                for tech in technical_skills:
                    if tech.lower() in token.lower() or token.lower() in tech.lower():
                        if token not in found_skills:
                            found_skills.append(token)
    except Exception as e:
        print(f"Error in NLTK processing: {e}")
    
    return list(set(found_skills))  # Remove duplicates

def extract_experience(text):
    """Extract work experience from resume text"""
    experience_keywords = ['experience', 'work history', 'employment', 'career', 'professional experience']
    education_keywords = ['education', 'academic', 'degree', 'university', 'college', 'school']
    
    lines = text.split('\n')
    experience_section = []
    in_experience = False
    
    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        
        # Check if we're entering experience section
        if any(keyword in line_lower for keyword in experience_keywords):
            in_experience = True
            continue
        
        # Check if we're leaving experience section
        if in_experience and any(keyword in line_lower for keyword in education_keywords + ['skills', 'projects', 'certifications']):
            break
        
        if in_experience and line.strip():
            # Look for job titles, companies, and dates
            if re.search(r'\d{4}', line) or any(word in line_lower for word in ['manager', 'developer', 'engineer', 'analyst', 'specialist', 'coordinator']):
                experience_section.append(line.strip())
    
    return experience_section

def extract_education(text):
    """Extract education information from resume text"""
    education_keywords = ['education', 'academic', 'degree', 'university', 'college', 'school']
    
    lines = text.split('\n')
    education_section = []
    in_education = False
    
    for line in lines:
        line_lower = line.lower().strip()
        
        if any(keyword in line_lower for keyword in education_keywords):
            in_education = True
            continue
        
        if in_education and line.strip():
            # Stop if we hit another section
            if any(keyword in line_lower for keyword in ['experience', 'skills', 'projects', 'certifications']):
                break
            
            # Look for degrees, institutions, and dates
            if any(word in line_lower for word in ['bachelor', 'master', 'phd', 'degree', 'university', 'college']) or re.search(r'\d{4}', line):
                education_section.append(line.strip())
    
    return education_section

def extract_keywords_simple(text):
    """Extract important keywords using simple frequency analysis"""
    try:
        # Remove common stop words
        stop_words = set(stopwords.words('english'))
        
        # Additional stop words specific to resumes
        additional_stop_words = {
            'resume', 'cv', 'curriculum', 'vitae', 'name', 'address', 'phone', 'email',
            'references', 'available', 'upon', 'request', 'page', 'www', 'http', 'https'
        }
        stop_words.update(additional_stop_words)
        
        # Clean text
        cleaned_text = re.sub(r'[^a-zA-Z\s]', '', text.lower())
        words = [word for word in cleaned_text.split() if word not in stop_words and len(word) > 2]
        
        # Count word frequency
        word_freq = Counter(words)
        
        # Get top keywords (appearing more than once)
        keywords = [word for word, freq in word_freq.most_common(30) if freq > 1]
        
        return keywords[:20]  # Return top 20
        
    except Exception as e:
        print(f"Error extracting keywords: {e}")
        return []

def identify_sections(text):
    """Identify which sections are present in the resume"""
    sections = {
        'contact': False,
        'summary': False,
        'experience': False,
        'education': False,
        'skills': False,
        'projects': False,
        'certifications': False
    }
    
    text_lower = text.lower()
    
    # Check for contact info
    if '@' in text or re.search(r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', text):
        sections['contact'] = True
    
    # Check for other sections
    section_keywords = {
        'summary': ['summary', 'objective', 'profile', 'about'],
        'experience': ['experience', 'work history', 'employment', 'career'],
        'education': ['education', 'academic', 'degree', 'university', 'college'],
        'skills': ['skills', 'technical skills', 'competencies', 'technologies'],
        'projects': ['projects', 'portfolio', 'work samples'],
        'certifications': ['certifications', 'certificates', 'licenses']
    }
    
    for section, keywords in section_keywords.items():
        if any(keyword in text_lower for keyword in keywords):
            sections[section] = True
    
    return sections

def calculate_simple_similarity(text1, text2):
    """Calculate simple text similarity using word overlap"""
    try:
        if not text1 or not text2:
            return 0.0
        
        # Tokenize and clean
        words1 = set(re.findall(r'\w+', text1.lower()))
        words2 = set(re.findall(r'\w+', text2.lower()))
        
        # Remove stop words
        stop_words = set(stopwords.words('english'))
        words1 = words1 - stop_words
        words2 = words2 - stop_words
        
        # Calculate Jaccard similarity
        intersection = len(words1 & words2)
        union = len(words1 | words2)
        
        if union == 0:
            return 0.0
        
        return (intersection / union) * 100
        
    except Exception as e:
        print(f"Error calculating similarity: {e}")
        return 0.0

def calculate_skill_match_score(user_skills, job_skills):
    """Calculate skill match score using simple string matching"""
    if not user_skills or not job_skills:
        return 0.0
    
    try:
        user_skills_lower = [skill.lower().strip() for skill in user_skills]
        job_skills_lower = [skill.lower().strip() for skill in job_skills]
        
        # Exact matches
        exact_matches = sum(1 for job_skill in job_skills_lower if job_skill in user_skills_lower)
        exact_match_score = (exact_matches / len(job_skills_lower)) if job_skills_lower else 0
        
        # Partial matches (contains)
        partial_matches = 0
        for job_skill in job_skills_lower:
            for user_skill in user_skills_lower:
                if job_skill in user_skill or user_skill in job_skill:
                    partial_matches += 1
                    break
        
        partial_match_score = (partial_matches / len(job_skills_lower)) if job_skills_lower else 0
        
        # Combine scores (weighted average)
        combined_score = (
            exact_match_score * 0.7 +  # 70% weight for exact matches
            partial_match_score * 0.3   # 30% weight for partial matches
        )
        
        return min(combined_score * 100, 100.0)
        
    except Exception as e:
        print(f"Error calculating skill match: {e}")
        # Fallback to simple string matching
        user_skills_lower = [skill.lower() for skill in user_skills]
        job_skills_lower = [skill.lower() for skill in job_skills]
        
        matches = sum(1 for job_skill in job_skills_lower if job_skill in user_skills_lower)
        return (matches / len(job_skills_lower)) * 100 if job_skills_lower else 0

def calculate_text_similarity(text1, text2):
    """Calculate similarity between two texts using simple word overlap"""
    return calculate_simple_similarity(text1, text2)

def enhanced_skill_extraction(text):
    """Enhanced skill extraction with better pattern matching"""
    
    # Comprehensive skill database with variations
    skill_patterns = {
        # Programming Languages
        'python': ['python', 'py'],
        'javascript': ['javascript', 'js', 'ecmascript'],
        'java': ['java'],
        'c++': ['c++', 'cpp', 'c plus plus'],
        'c#': ['c#', 'csharp', 'c sharp'],
        'php': ['php'],
        'ruby': ['ruby'],
        'go': ['golang', 'go'],
        'rust': ['rust'],
        'swift': ['swift'],
        'kotlin': ['kotlin'],
        'typescript': ['typescript', 'ts'],
        'scala': ['scala'],
        'r': ['r programming', 'r language'],
        
        # Web Technologies
        'html': ['html', 'html5'],
        'css': ['css', 'css3'],
        'react': ['react', 'reactjs', 'react.js'],
        'angular': ['angular', 'angularjs'],
        'vue': ['vue', 'vuejs', 'vue.js'],
        'node.js': ['nodejs', 'node.js', 'node js'],
        'express': ['express', 'expressjs', 'express.js'],
        'django': ['django'],
        'flask': ['flask'],
        'spring': ['spring', 'spring boot'],
        'laravel': ['laravel'],
        'rails': ['rails', 'ruby on rails'],
        
        # Databases
        'mysql': ['mysql'],
        'postgresql': ['postgresql', 'postgres'],
        'mongodb': ['mongodb', 'mongo'],
        'redis': ['redis'],
        'sqlite': ['sqlite'],
        'oracle': ['oracle', 'oracle db'],
        
        # Cloud & DevOps
        'aws': ['aws', 'amazon web services'],
        'azure': ['azure', 'microsoft azure'],
        'gcp': ['gcp', 'google cloud', 'google cloud platform'],
        'docker': ['docker'],
        'kubernetes': ['kubernetes', 'k8s'],
        'jenkins': ['jenkins'],
        'git': ['git'],
        
        # Data Science
        'machine learning': ['machine learning', 'ml'],
        'deep learning': ['deep learning', 'dl'],
        'tensorflow': ['tensorflow'],
        'pytorch': ['pytorch'],
        'pandas': ['pandas'],
        'numpy': ['numpy'],
        'scikit-learn': ['scikit-learn', 'sklearn'],
    }
    
    found_skills = set()
    text_lower = text.lower()
    
    # Extract skills using patterns
    for skill, patterns in skill_patterns.items():
        for pattern in patterns:
            if pattern in text_lower:
                found_skills.add(skill.title())
                break
    
    # Additional pattern-based extraction
    # Look for common skill sections
    skill_section_pattern = r'(?i)(?:skills?|technologies?|competencies)[\s\n]*[:\-]?\s*([^\n]+(?:\n[^\n]+)*?)(?=\n\s*(?:[A-Z][^:\n]*:|$))'
    skill_matches = re.findall(skill_section_pattern, text)
    
    for match in skill_matches:
        # Split by common delimiters
        skills_in_section = re.split(r'[,;|\n•\-\*]', match)
        for skill in skills_in_section:
            skill = skill.strip()
            if len(skill) > 2 and skill.lower() not in ['and', 'or', 'with', 'including']:
                found_skills.add(skill.title())
    
    return list(found_skills)

def analyze_resume(resume, job_id=None):
    """Analyze resume and provide scoring and recommendations"""
    
    # Basic analysis
    overall_score = calculate_overall_score(resume)
    ats_score = calculate_ats_score(resume)
    strengths = identify_strengths(resume)
    improvements = identify_improvements(resume)
    missing_keywords = []
    
    # Job-specific analysis if job_id provided
    if job_id:
        try:
            job = Job.objects.get(id=job_id)
            missing_keywords = find_missing_keywords(resume, job)
        except Job.DoesNotExist:
            pass
    
    # Create analysis record
    analysis = ResumeAnalysis.objects.create(
        resume=resume,
        job_id=job_id if job_id else None,
        overall_score=overall_score,
        ats_score=ats_score,
        strengths=strengths,
        improvements=improvements,
        missing_keywords=missing_keywords,
        section_scores=resume.parsed_data.get('sections', {})
    )
    
    return analysis

def calculate_overall_score(resume):
    """Calculate overall resume score based on multiple factors"""
    score = 0
    
    # Section completeness (40 points)
    sections = resume.parsed_data.get('sections', {})
    section_weights = {
        'contact': 10,
        'experience': 15,
        'education': 10,
        'skills': 5
    }
    
    for section, weight in section_weights.items():
        if sections.get(section, False):
            score += weight
    
    # Content quality (30 points)
    skills = resume.parsed_data.get('skills', [])
    if len(skills) >= 10:
        score += 20
    elif len(skills) >= 5:
        score += 15
    elif len(skills) >= 3:
        score += 10
    elif len(skills) >= 1:
        score += 5
    
    experience = resume.parsed_data.get('experience', [])
    if len(experience) >= 3:
        score += 10
    elif len(experience) >= 1:
        score += 5
    
    # Text quality and structure (30 points)
    text_length = len(resume.extracted_text)
    if text_length >= 2000:
        score += 15
    elif text_length >= 1000:
        score += 12
    elif text_length >= 500:
        score += 8
    elif text_length >= 200:
        score += 5
    
    # Keywords and relevance
    keywords = resume.parsed_data.get('keywords', [])
    if len(keywords) >= 15:
        score += 10
    elif len(keywords) >= 10:
        score += 7
    elif len(keywords) >= 5:
        score += 5
    
    # Contact info bonus
    contact_info = resume.parsed_data.get('contact_info', {})
    if contact_info.get('email'):
        score += 3
    if contact_info.get('phone'):
        score += 2
    
    return min(score, 100)  # Cap at 100

def calculate_ats_score(resume):
    """Calculate ATS compatibility score"""
    score = 0
    
    # File format (20 points)
    if resume.original_filename.lower().endswith(('.pdf', '.docx')):
        score += 20
    
    # Section headers (30 points)
    sections = resume.parsed_data.get('sections', {})
    required_sections = ['contact', 'experience', 'skills']
    for section in required_sections:
        if sections.get(section):
            score += 10
    
    # Skills and keywords (30 points)
    skills = resume.parsed_data.get('skills', [])
    if len(skills) >= 8:
        score += 30
    elif len(skills) >= 5:
        score += 20
    elif len(skills) >= 3:
        score += 15
    elif len(skills) >= 1:
        score += 10
    
    # Text structure and readability (20 points)
    text = resume.extracted_text
    lines = text.split('\n')
    non_empty_lines = [line for line in lines if line.strip()]
    
    if len(non_empty_lines) >= 20:  # Good structure
        score += 20
    elif len(non_empty_lines) >= 15:
        score += 15
    elif len(non_empty_lines) >= 10:
        score += 10
    elif len(non_empty_lines) >= 5:
        score += 5
    
    return min(score, 100)

def identify_strengths(resume):
    """Identify resume strengths"""
    strengths = []
    
    sections = resume.parsed_data.get('sections', {})
    skills = resume.parsed_data.get('skills', [])
    contact_info = resume.parsed_data.get('contact_info', {})
    experience = resume.parsed_data.get('experience', [])
    
    if sections.get('contact') and contact_info:
        strengths.append("Complete contact information provided")
    
    if len(skills) >= 10:
        strengths.append("Comprehensive technical skills section")
    elif len(skills) >= 5:
        strengths.append("Good technical skills coverage")
    
    if sections.get('experience') and len(experience) >= 2:
        strengths.append("Detailed work experience documented")
    
    if sections.get('education'):
        strengths.append("Educational background clearly presented")
    
    if len(resume.extracted_text) >= 1500:
        strengths.append("Comprehensive resume content")
    
    if contact_info.get('linkedin'):
        strengths.append("Professional LinkedIn profile included")
    
    if sections.get('projects'):
        strengths.append("Project experience highlighted")
    
    if sections.get('certifications'):
        strengths.append("Professional certifications listed")
    
    return strengths

def identify_improvements(resume):
    """Identify areas for improvement"""
    improvements = []
    
    sections = resume.parsed_data.get('sections', {})
    skills = resume.parsed_data.get('skills', [])
    contact_info = resume.parsed_data.get('contact_info', {})
    
    if not sections.get('summary'):
        improvements.append("Add a professional summary or objective statement")
    
    if len(skills) < 5:
        improvements.append("Include more relevant technical and soft skills")
    
    if not sections.get('experience'):
        improvements.append("Add detailed work experience section")
    
    if len(resume.extracted_text) < 800:
        improvements.append("Expand resume content with more specific details and achievements")
    
    if not contact_info.get('linkedin'):
        improvements.append("Include LinkedIn profile URL")
    
    if not sections.get('projects'):
        improvements.append("Consider adding a projects section to showcase your work")
    
    if not sections.get('certifications'):
        improvements.append("Add relevant certifications or professional development")
    
    # Check for quantifiable achievements
    if not re.search(r'\d+%|\$\d+|\d+\s*(years?|months?)', resume.extracted_text):
        improvements.append("Include quantifiable achievements and metrics")
    
    return improvements

def find_missing_keywords(resume, job):
    """Find keywords missing from resume compared to job requirements"""
    resume_text = resume.extracted_text.lower()
    job_skills = [skill.lower() for skill in job.skills] if job.skills else []
    
    # Extract keywords from job description
    job_keywords = extract_keywords_simple(job.description)
    
    missing_keywords = []
    
    # Check missing skills
    for skill in job_skills:
        if skill not in resume_text:
            missing_keywords.append(skill.title())
    
    # Check missing job description keywords
    for keyword in job_keywords[:10]:  # Top 10 keywords
        if keyword not in resume_text and keyword.lower() not in missing_keywords:
            missing_keywords.append(keyword.title())
    
    return list(set(missing_keywords))[:15]  # Return top 15 unique missing keywords